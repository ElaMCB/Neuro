"""
Resume Preparation and Tailoring System
Prepares resumes customized for specific job postings
"""

import os
import re
import json
from typing import Dict, List, Optional
from pathlib import Path
from datetime import datetime

# Try to import Word document support
try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

class ResumePreparer:
    """Prepares and tailors resumes for specific job applications"""
    
    def __init__(self, base_resume_path: Optional[str] = None):
        self.base_resume_path = base_resume_path
        self.templates_dir = "resume_templates"
        Path(self.templates_dir).mkdir(exist_ok=True)
        
    def prepare_for_job(self, job_title: str, company: str, job_description: str, 
                       job_url: str, base_resume_content: Optional[str] = None) -> Dict[str, str]:
        """Prepare a tailored resume for a specific job"""
        
        print(f"\n📝 Preparing resume for {job_title} at {company}...")
        
        # Load base resume
        if base_resume_content is None:
            base_resume_content = self._load_base_resume()
        
        if not base_resume_content:
            print("⚠️  No base resume found. Please provide your resume.")
            return {}
        
        # Extract keywords from job description
        keywords = self._extract_keywords(job_description)
        required_skills = self._extract_required_skills(job_description)
        
        # Tailor resume
        tailored_resume = self._tailor_resume(
            base_resume_content,
            job_title,
            company,
            keywords,
            required_skills,
            job_description
        )
        
        # Generate cover letter
        cover_letter = self._generate_cover_letter(
            job_title,
            company,
            job_description,
            required_skills
        )
        
        # Save tailored documents
        timestamp = datetime.now().strftime("%Y%m%d")
        safe_company = re.sub(r'[^\w\s-]', '', company).replace(' ', '_').lower()
        safe_title = re.sub(r'[^\w\s-]', '', job_title).replace(' ', '_').lower()
        
        resume_filename = f"{self.templates_dir}/resume_{safe_company}_{safe_title}_{timestamp}.txt"
        cover_filename = f"{self.templates_dir}/cover_{safe_company}_{safe_title}_{timestamp}.txt"
        
        self._save_file(resume_filename, tailored_resume)
        self._save_file(cover_filename, cover_letter)
        
        print(f"✅ Resume tailored and saved to: {resume_filename}")
        print(f"✅ Cover letter saved to: {cover_filename}")
        
        return {
            "resume_path": resume_filename,
            "cover_letter_path": cover_filename,
            "keywords": keywords,
            "required_skills": required_skills,
            "job_url": job_url
        }
    
    def _load_base_resume(self) -> Optional[str]:
        """Load base resume from file (supports .txt, .md, .docx)"""
        # Try multiple possible locations
        possible_paths = [
            self.base_resume_path,
            "examples/resume/outputs/resume_prompt_engineer_short.txt",
            "examples/resume/outputs/resume_prompt_engineer_keywords.txt",
            "resume.txt",
            "my_resume.txt",
            "resume.docx",
            "my_resume.docx"
        ]
        
        for path in possible_paths:
            if path and os.path.exists(path):
                try:
                    # Check file extension
                    file_ext = os.path.splitext(path)[1].lower()
                    
                    if file_ext == '.docx':
                        # Load Word document
                        if not DOCX_SUPPORT:
                            print(f"⚠️  Word document support not available. Install: pip install python-docx")
                            print(f"   Trying to read as text file...")
                            # Fallback: try reading as text
                            with open(path, 'r', encoding='utf-8', errors='ignore') as f:
                                return f.read()
                        else:
                            return self._read_docx(path)
                    else:
                        # Read as text file
                        with open(path, 'r', encoding='utf-8') as f:
                            return f.read()
                except Exception as e:
                    print(f"⚠️  Error reading {path}: {e}")
                    continue
        
        return None
    
    def _read_docx(self, path: str) -> str:
        """Read content from a Word document (.docx)"""
        try:
            doc = Document(path)
            text_content = []
            
            # Extract text from all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
        except Exception as e:
            print(f"⚠️  Error reading Word document {path}: {e}")
            raise
    
    def _extract_keywords(self, job_description: str) -> List[str]:
        """Extract important keywords from job description"""
        # Common AI/ML keywords
        ai_keywords = [
            "python", "pytorch", "tensorflow", "llm", "gpt", "transformer",
            "nlp", "machine learning", "deep learning", "neural network",
            "prompt engineering", "fine-tuning", "rag", "embeddings",
            "api", "docker", "kubernetes", "aws", "mlops", "ci/cd",
            "git", "github", "research", "experimentation", "evaluation"
        ]
        
        description_lower = job_description.lower()
        found_keywords = [kw for kw in ai_keywords if kw in description_lower]
        
        return found_keywords
    
    def _extract_required_skills(self, job_description: str) -> List[str]:
        """Extract required skills from job description"""
        # Look for "required", "must have", "experience with" patterns
        required_patterns = [
            r"required[:\s]+(.+?)(?:\.|;|\n|preferred)",
            r"must have[:\s]+(.+?)(?:\.|;|\n|nice to have)",
            r"experience with[:\s]+(.+?)(?:\.|;|\n)",
        ]
        
        skills = []
        for pattern in required_patterns:
            matches = re.finditer(pattern, job_description, re.IGNORECASE)
            for match in matches:
                skill_text = match.group(1)
                # Extract individual skills (Python, PyTorch, etc.)
                individual_skills = re.findall(r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)\b', skill_text)
                skills.extend(individual_skills)
        
        return list(set(skills))[:10]  # Top 10 unique skills
    
    def _tailor_resume(self, base_resume: str, job_title: str, company: str,
                      keywords: List[str], required_skills: List[str],
                      job_description: str) -> str:
        """Tailor the base resume for the specific job"""
        
        # Add tailored summary section if needed
        tailored_summary = self._create_tailored_summary(
            job_title,
            company,
            keywords,
            required_skills
        )
        
        # Enhance skills section to match job requirements
        enhanced_resume = self._enhance_skills_section(
            base_resume,
            keywords,
            required_skills
        )
        
        # Ensure relevant keywords appear in resume
        enhanced_resume = self._ensure_keyword_coverage(
            enhanced_resume,
            keywords
        )
        
        # Add tailored header
        header = f"""
RESUME TAILORED FOR: {job_title}
COMPANY: {company}
DATE: {datetime.now().strftime('%Y-%m-%d')}

{tailored_summary}

{'=' * 60}

"""
        
        return header + enhanced_resume
    
    def _create_tailored_summary(self, job_title: str, company: str,
                                 keywords: List[str], required_skills: List[str]) -> str:
        """Create a tailored professional summary"""
        
        # Build summary emphasizing matching skills
        matching_tech = [k for k in keywords if k in ['python', 'pytorch', 'llm', 'gpt', 'transformer', 'nlp']]
        
        summary = f"""
TAILORED PROFESSIONAL SUMMARY
-----------------------------
Target Position: {job_title} at {company}

Key Alignment:
"""
        
        if matching_tech:
            summary += f"• Strong experience with {', '.join(matching_tech[:3])}\n"
        
        if 'prompt engineer' in job_title.lower() or 'prompt' in str(required_skills).lower():
            summary += "• Creator of Neuro programming language - demonstrating intent-driven AI development\n"
            summary += "• Experience with prompt engineering and LLM optimization\n"
        
        if 'junior' in job_title.lower() or 'entry' in job_title.lower():
            summary += "• Eager to grow in AI engineering with strong foundation in Python and AI workflows\n"
        else:
            summary += "• Transitioning from QA leadership to AI engineering with hands-on project experience\n"
        
        summary += "\nNOTE: This resume has been tailored to highlight relevant experience for this position.\n"
        
        return summary
    
    def _enhance_skills_section(self, resume: str, keywords: List[str],
                               required_skills: List[str]) -> str:
        """Enhance the skills section to match job requirements"""
        
        # Find skills section
        skills_section_pattern = r'(TECHNICAL SKILLS|SKILLS|TECHNICAL COMPETENCIES)[:|\s]*\n(.+?)(?=\n[A-Z]|\n\n|$)'
        
        match = re.search(skills_section_pattern, resume, re.IGNORECASE | re.DOTALL)
        if match:
            skills_section = match.group(2)
            
            # Add missing required skills if not present
            skills_lower = skills_section.lower()
            missing_skills = [s for s in required_skills if s.lower() not in skills_lower]
            
            if missing_skills:
                # Add a note about additional relevant skills
                enhancement = f"\nADDITIONAL RELEVANT SKILLS FOR THIS POSITION:\n"
                for skill in missing_skills[:5]:
                    enhancement += f"• {skill} (learning/familiar)\n"
                
                resume = resume.replace(match.group(0), match.group(0) + enhancement)
        
        return resume
    
    def _ensure_keyword_coverage(self, resume: str, keywords: List[str]) -> str:
        """Ensure important keywords appear in resume (for ATS systems)"""
        
        resume_lower = resume.lower()
        missing_keywords = [kw for kw in keywords[:5] if kw not in resume_lower]
        
        if missing_keywords:
            # Add a relevant skills note
            note = f"\n\nRELEVANT KEYWORDS FOR THIS POSITION: {', '.join(missing_keywords)}\n"
            note += "(These technologies align with the role's requirements)\n"
            resume += note
        
        return resume
    
    def _generate_cover_letter(self, job_title: str, company: str,
                              job_description: str, required_skills: List[str]) -> str:
        """Generate a tailored cover letter"""
        
        # Extract company focus from description
        company_focus = self._extract_company_focus(job_description)
        
        cover_letter = f"""
COVER LETTER
{'=' * 60}

Date: {datetime.now().strftime('%B %d, %Y')}

Hiring Manager
{company}

Subject: Application for {job_title} Position

Dear Hiring Manager,

I am writing to express my strong interest in the {job_title} position at {company}. 
"""
        
        # Add tailored paragraph based on company focus
        if company_focus:
            cover_letter += f"""
{company} is clearly focused on {company_focus}, which aligns perfectly with my background in 
AI system development and my recent work exploring intent-driven programming concepts through 
projects like Neuro.
"""
        else:
            cover_letter += f"""
{company}'s mission aligns with my background in AI system development and my recent work 
exploring intent-driven programming concepts through projects like Neuro.
"""
        
        # Add skills paragraph
        if required_skills:
            top_skills = required_skills[:3]
            cover_letter += f"""
I bring relevant experience with {', '.join(top_skills)} and am particularly excited about 
the opportunity to contribute to {company}'s AI initiatives.
"""
        
        cover_letter += f"""
My background includes:
"""
        
        # Add relevant experience points
        cover_letter += """
• Creating Neuro, an intent-driven programming language exploration that demonstrates practical 
  AI engineering concepts including natural language understanding and automated workflow generation

• Strong foundation in Python, AI workflows, and system architecture from QA leadership roles

• Passion for making AI development more accessible and building tools that developers can 
  actually use

"""
        
        if 'prompt' in job_title.lower() or 'prompt' in job_description.lower():
            cover_letter += """
• Hands-on experience with prompt engineering through developing Neuro's natural language 
  understanding capabilities and working with LLM systems

"""
        
        cover_letter += f"""
I am excited about the opportunity to bring my AI engineering experience and enthusiasm to 
{company}. I would welcome the chance to discuss how my background in AI systems could 
contribute to your team.

Thank you for considering my application.

Best regards,
[Your Name]
[Your Email]
[Your Phone]
[LinkedIn/GitHub URLs]
"""
        
        return cover_letter
    
    def _extract_company_focus(self, job_description: str) -> str:
        """Extract what the company focuses on from job description"""
        # Look for company mission/focus statements
        focus_patterns = [
            r"mission is to (.+?)(?:\.|,|and)",
            r"focused on (.+?)(?:\.|,|and)",
            r"building (.+?)(?:\.|,|to)",
        ]
        
        for pattern in focus_patterns:
            match = re.search(pattern, job_description, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        
        return ""
    
    def _save_file(self, filepath: str, content: str):
        """Save content to file"""
        Path(filepath).parent.mkdir(parents=True, exist_ok=True)
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)

